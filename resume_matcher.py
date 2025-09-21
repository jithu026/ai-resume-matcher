# resume_matcher.py
"""
Gemini-ready resume_matcher.py

Usage:
- Set GEMINI_API_KEY env var to use Google Gemini (AI Studio) model.
- If GEMINI_API_KEY not set, the module will run in MOCK mode (safe for development).
- Exposes optimize_resume(job_description, resume_text) -> dict
"""

import os
import re
import json
import io
from typing import List, Dict
from docx import Document

# Try to import Gemini SDK. If not available, we'll fall back to mock.
USE_GEMINI = False
try:
    import google.generativeai as genai  # type: ignore
    USE_GEMINI = True
except Exception:
    USE_GEMINI = False

GEMINI_KEY = os.getenv("AIzaSyD9nLqtNRf4h3VGz8nUQU1Vv1yHiCS6_0g")
MODEL = os.getenv("GEMINI_MODEL", "gemini-1.5-flash")  # change if you have access to different model

# Determine runtime mode
USE_REAL_GEMINI = bool(GEMINI_KEY) and USE_GEMINI

if USE_REAL_GEMINI:
    # Configure Gemini client
    genai.configure(api_key=GEMINI_KEY)
    # If your installed SDK expects different usage, adjust below accordingly.
    # We will use a thin wrapper call_llm() that uses the installed SDK.
else:
    # Running in mock mode (no external API calls)
    pass


def call_llm(prompt: str, system: str = None, temperature: float = 0.3, max_tokens: int = 500) -> str:
    """
    Unified LLM call function.
    - If GEMINI available and GEMINI_KEY set, call Gemini.
    - Otherwise return heuristic/mock output for development.
    """
    if USE_REAL_GEMINI:
        try:
            # Using model.generate_content style if available
            # This is a general approach — if your genai version differs, adapt to genai.generate_text(...)
            model = genai.GenerativeModel(MODEL)
            response = model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=temperature,
                    max_output_tokens=max_tokens,
                )
            )
            # response.text may be the main content
            text = getattr(response, "text", None)
            if not text:
                # Some SDK variants provide a different structure
                text = json.dumps(response)
            return text.strip()
        except Exception:
            # If that fails, try alternative (older) API:
            try:
                resp = genai.generate_text(model=MODEL, prompt=prompt, temperature=temperature, max_output_tokens=max_tokens)
                return getattr(resp, "text", str(resp)).strip()
            except Exception:
                # fallback to mock
                return _mock_call_llm(prompt)
    else:
        return _mock_call_llm(prompt)


def _mock_call_llm(prompt: str) -> str:
    """
    Simple mock for dev when no API key present. Returns reasonable placeholders.
    """
    import re, json
    # If the prompt asks for a JSON array of top keywords:
    if "Extract the top" in prompt and "JSON array" in prompt:
        common = ["AWS", "Docker", "Kubernetes", "CI/CD", "Python", "Azure", "GCP", "Linux"]
        found = []
        for c in common:
            if re.search(r"\b" + re.escape(c) + r"\b", prompt, flags=re.IGNORECASE):
                found.append(c)
        if not found:
            # fallback: collect capitalized tokens
            toks = re.findall(r"\b[A-Z][A-Za-z0-9+#\-_/.]{1,}\b", prompt)
            found = list(dict.fromkeys(toks))[:8]
        return json.dumps(found)

    # If the prompt requests a resume rewrite in markdown:
    if "Rewrite a concise \"Summary\"" in prompt or "### Summary" in prompt:
        m = re.search(r"Matched keywords:\s*\[(.*?)\]", prompt, re.S)
        mm = re.search(r"Missing keywords:\s*\[(.*?)\]", prompt, re.S)
        matched = []
        missing = []
        if m:
            matched = [x.strip().strip('"').strip("'") for x in m.group(1).split(",") if x.strip()]
        if mm:
            missing = [x.strip().strip('"').strip("'") for x in mm.group(1).split(",") if x.strip()]

        summary = "### Summary\n"
        if matched:
            summary += f"Experienced professional with expertise in {', '.join(matched[:4])} and strong automation skills.\n\n"
        else:
            summary += "Experienced professional with technical and automation skills, focused on improving deployment workflows.\n\n"

        skills_lines = "### Skills\n"
        for s in matched:
            skills_lines += f"- {s}\n"
        for s in missing:
            skills_lines += f"- Familiar with {s}\n"

        bullets = "\n### Suggested bullets\n"
        bullets += "- Improved deployment automation through scripting and CI/CD processes.\n"
        bullets += "- Built automation scripts to reduce manual tasks and speed up deployments.\n"

        return summary + skills_lines + bullets

    # Default fallback
    return "[]"


def extract_keywords(job_description: str, top_n: int = 12) -> List[str]:
    """
    Extract top technical skills/keywords from JD.
    We ask the LLM to return a JSON array; fallback to heuristic extraction if parsing fails.
    """
    prompt = (
        f"Extract the top {top_n} technical skills and keywords from the following job description.\n"
        "Return a JSON array only, for example: [\"AWS\", \"Docker\", \"Kubernetes\"]\n\n"
        f"Job Description:\n{job_description}"
    )
    out = call_llm(prompt, temperature=0.0, max_tokens=300)
    # Try to parse JSON
    try:
        kw = json.loads(out)
        if isinstance(kw, list) and kw:
            return [str(k).strip() for k in kw if str(k).strip()]
    except Exception:
        pass

    # Heuristic fallback
    tokens = re.findall(r"[A-Za-z0-9+_\-/\.#]+", job_description)
    freq = {}
    for t in tokens:
        tl = t.lower()
        if len(tl) <= 1:
            continue
        freq[tl] = freq.get(tl, 0) + 1
    sorted_tokens = sorted(freq.items(), key=lambda x: -x[1])
    # Return top_n tokens (capitalized)
    return [t[0] for t in sorted_tokens[:top_n]]


def compute_match_score(keywords: List[str], resume_text: str) -> Dict:
    """
    Compute a simple keyword overlap based match score (0-100) and list matched/missing.
    """
    resume_low = resume_text.lower()
    matched = []
    missing = []
    for k in keywords:
        k_low = k.lower()
        parts = re.findall(r"\w+", k_low)
        if all(part in resume_low for part in parts if part):
            matched.append(k)
        else:
            missing.append(k)
    score = int(100 * len(matched) / max(1, len(keywords)))
    return {"score": score, "matched": matched, "missing": missing}


def generate_optimized_resume(job_description: str, resume_text: str, matched: List[str], missing: List[str]) -> str:
    """
    Ask LLM to rewrite Summary, Skills section and suggest achievement-style bullets.
    Returns Markdown text.
    """
    prompt = f"""
You are an expert resume writer. Given the Job Description and the Candidate Resume below, do the following:

1) Rewrite a concise "Summary" (1-2 sentences) that highlights the candidate's matching skills (use matched list).
2) Produce a "Skills" section as bullet list focusing on matched skills first, then phrasing for missing skills as "Familiar with X" or "Exposure to Y" if the resume doesn't explicitly contain them.
3) Suggest up to 3 concrete resume bullet(s) (achievement-style, metric if possible) that the candidate can use or adapt — do NOT make up false work history; only rephrase existing experience truthfully or suggest reasonable "learning" phrasing.

Provide the result in Markdown with headings "### Summary", "### Skills", "### Suggested bullets".

Job Description:
{job_description}

Candidate Resume:
{resume_text}

Matched keywords: {matched}
Missing keywords: {missing}
"""
    out = call_llm(prompt, temperature=0.25, max_tokens=600)
    return out


def make_docx_from_markdown(markdown_text: str) -> bytes:
    """
    Simple markdown -> docx conversion for the generated sections.
    """
    doc = Document()
    for line in markdown_text.splitlines():
        if line.startswith("### "):
            doc.add_heading(line.replace("### ", "").strip(), level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line.strip())
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def optimize_resume(job_description: str, resume_text: str) -> Dict:
    """
    High-level convenience function to run the full pipeline and return structured result.
    """
    keywords = extract_keywords(job_description)
    match = compute_match_score(keywords, resume_text)
    optimized_md = generate_optimized_resume(job_description, resume_text, match["matched"], match["missing"])
    docx_bytes = make_docx_from_markdown(optimized_md)
    return {
        "keywords": keywords,
        "score": match["score"],
        "matched": match["matched"],
        "missing": match["missing"],
        "optimized_markdown": optimized_md,
        "optimized_docx_bytes": docx_bytes
    }


# CLI quick test
if __name__ == "__main__":
    print("Gemini available:", USE_GEMINI)
    print("GEMINI_KEY set:", bool(GEMINI_KEY))
    print("Running in REAL GEMINI mode:", USE_REAL_GEMINI)
    print("----\nPaste Job Description (end with a single line 'END'):\n")
    jd_lines = []
    while True:
        try:
            line = input()
        except EOFError:
            break
        if line.strip() == "END":
            break
        jd_lines.append(line)
    jd = "\n".join(jd_lines).strip()
    if not jd:
        print("No JD provided — exiting.")
        raise SystemExit(0)

    print("\nPaste Resume text (end with a single line 'END'):\n")
    res_lines = []
    while True:
        try:
            line = input()
        except EOFError:
            break
        if line.strip() == "END":
            break
        res_lines.append(line)
    resume_txt = "\n".join(res_lines).strip()

    result = optimize_resume(jd, resume_txt)
    print("\n=== ATS-like Match Score ===")
    print(result["score"])
    print("\n=== Extracted Keywords ===")
    print(result["keywords"])
    print("\n=== Matched ===")
    print(result["matched"])
    print("\n=== Missing ===")
    print(result["missing"])
    print("\n=== Optimized Markdown ===\n")
    print(result["optimized_markdown"][:2000])  # print first chunk
